import customtkinter as ctk
from PIL import Image, ImageTk
import tkinter as tk
from tkinter import filedialog
import requests
import json
from dotenv import load_dotenv
import os
import csv
from openpyxl import load_workbook
import shutil
from openpyxl.styles import PatternFill
from docx import Document
from datetime import datetime
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from geopy.geocoders import Nominatim
from geopy.distance import geodesic


load_dotenv()

API_KEY = os.getenv("API_KEY")
BOARD_ID = os.getenv("BOARD_ID")
API_URL = os.getenv("API_URL")

# GraphQL query with corrected types
query = """
query($boardId: [ID!], $cursor: String) {
  boards(ids: $boardId) {
    items_page(limit: 100, cursor: $cursor) {
      cursor
      items {
        name
        column_values {
          column {
            title
          }
          text
        }
      }
    }
  }
}
"""

# Update saved API key
def update_api_key():
    global API_KEY
    file_path = ".env"
    new_key = entry.get()
    if new_key:  
        with open(file_path, 'r') as file:
            lines = file.readlines()
        with open(file_path, 'w') as file:
            for line in lines:
                if line.startswith("API_KEY"):
                    file.write(f'{"API_KEY"}={new_key}\n')
                else:
                    file.write(line)
        API_KEY = new_key
        tk.messagebox.showinfo("Succès", "API Key a été mise à jour avec succès!")
    else:
        tk.messagebox.showerror("Erreur", "API Key ne peut pas être vide.")

# Update saved Board ID
def update_board_id():
    global BOARD_ID
    file_path = ".env"
    new_key = boardEntry.get()
    if new_key:  
        with open(file_path, 'r') as file:
            lines = file.readlines()
        with open(file_path, 'w') as file:
            for line in lines:
                if line.startswith("BOARD_ID"):
                    file.write(f'{"BOARD_ID"}={new_key}\n')
                else:
                    file.write(line)
        BOARD_ID = new_key
        tk.messagebox.showinfo("Succès", "Board ID a été mise à jour avec succès!")
    else:
        tk.messagebox.showerror("Erreur", "Board ID ne peut pas être vide.")

def toggle_switch_excel():
    if switchExcel.get() == 1:
        chooseBtn.grid(row=0, column=1, sticky="e", pady=10, padx=10)
    else:
        excelPath.delete(0, tk.END)
        excelFilePath.configure(text="")
        chooseBtn.grid_forget()

def toggle_switch_pdf():
    if switchPDF.get() == 1:
        chooseBtnPDF.grid(row=0, column=1, sticky="e", pady=10, padx=10)
    else:
        pdfPath.delete(0, tk.END)
        pdfFilePath.configure(text="")
        chooseBtnPDF.grid_forget()

# Open File Explorer to choose path
def choose_folder(typeExtract):
    folder_path = filedialog.askdirectory()
    if typeExtract == "excel":
        excelPath.delete(0, tk.END)
        excelPath.insert(0, folder_path)
        excelFilePath.configure(text="*" + folder_path)
    elif typeExtract == "pdf":
        pdfPath.delete(0, tk.END)
        pdfPath.insert(0, folder_path)
        pdfFilePath.configure(text="*" + folder_path)

# Export Button
def export_info():
    if switchExcel.get() == 0 and switchPDF.get() == 0:
        tk.messagebox.showerror("Erreur", "Vous devez sélectionner une méthode d'exportation.")
    else:
        if switchExcel.get() == 1:
            excelWorks = export_excel()
            if excelWorks:
                tk.messagebox.showinfo("Succès", "Exportation Excel réussie!")
        if switchPDF.get() == 1:
            pdfWorks = export_pdf()
            if pdfWorks:
                tk.messagebox.showinfo("Succès", "Exportation Word réussie!")


# Export Word
def export_pdf():
    path = pdfPath.get()
    if path:
        all_items = call_api()
        reformatted_items = reformat_items(all_items)
        data = json.loads(json.dumps(reformatted_items))

        doc = Document('./template.docx')

        geolocator = Nominatim(user_agent="geo_distance")

        for para in doc.paragraphs:
            if "Projet Complet d’ExperiSens" in para.text:
                for i in range(len(data)):
                    if data[i]["column_values"]["Statut"] == "Complété":
                        project = ""
                        if data[i]["column_values"]["Code Projet"] != "":
                            project += f'{data[i]["column_values"]["Code Projet"]}'

                        location_name = data[i]["column_values"]["Lieu"]
                        if location_name != "":
                            location = geolocator.geocode(location_name)
            
                            if location is None:
                                print(f"Could not geocode location: {location_name}")
                                continue

                            # Define locations
                            ITHQ = geolocator.geocode("ITHQ, Montreal, QC, Canada")

                            #Coords
                            coords_ithq = (ITHQ.latitude, ITHQ.longitude)
                            coords = (location.latitude, location.longitude)

                            distance = geodesic(coords_ithq, coords).kilometers

                            run = para.add_run(f"\n {project} {data[i]['name'].capitalize()} {distance:.2f} km.")
                            run.font.size = Pt(9)                          
                            if distance > 100:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        else:
                            run = para.add_run(f"\n {project} {data[i]['name'].capitalize()}")
                            run.font.size = Pt(9)
            if "Projet Non Complet d'ExperiSens" in para.text:
                for i in range(len(data)):
                    if data[i]["column_values"]["Statut"] != "Complété":
                        project = ""
                        if data[i]["column_values"]["Code Projet"] != "":
                            project += f'{data[i]["column_values"]["Code Projet"]}'
                        
                        location_name = data[i]["column_values"]["Lieu"]
                        if location_name != "":
                            location = geolocator.geocode(location_name)
            
                            if location is None:
                                print(f"Could not geocode location: {location_name}")
                                continue

                            # Define locations
                            ITHQ = geolocator.geocode("ITHQ, Montreal, QC, Canada")

                            #Coords
                            coords_ithq = (ITHQ.latitude, ITHQ.longitude)
                            coords = (location.latitude, location.longitude)

                            distance = geodesic(coords_ithq, coords).kilometers

                            run = para.add_run(f"\n {project} {data[i]['name'].capitalize()} {distance:.2f} km.")
                            run.font.size = Pt(9)                          
                            if distance > 100:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        else:
                            run = para.add_run(f"\n {project} {data[i]['name'].capitalize()}")
                            run.font.size = Pt(9)

        doc.save(f"{path}/{datetime.now().year}-{datetime.now().year + 1} Requête annuelle.docx")

        return True
    else:
        tk.messagebox.showerror("Erreur", "Vous devez sélectionner un dossier pour enregistrer.")
    return False



# Export Excel
def export_excel():
    path = excelPath.get()
    if path:
        all_items = call_api()
        reformatted_items = reformat_items(all_items)
        #print(json.dumps(reformatted_items, indent=2))
        #print("-------")
        data = json.loads(json.dumps(reformatted_items))

        excel_template = './template.xlsx'
        workbook_path = path + "/projets.xlsx"

        shutil.copyfile(excel_template, workbook_path)

        workbook = load_workbook(workbook_path)

        sheet = workbook['ExperiSens']

        # Cell colours

        # Statut
        afaire = PatternFill(start_color="7b7f90", end_color="7b7f90", fill_type="solid")
        planifie = PatternFill(start_color="639cc9", end_color="639cc9", fill_type="solid")
        encours = PatternFill(start_color="edbf72", end_color="edbf72", fill_type="solid")
        bloque = PatternFill(start_color="d0717f", end_color="d0717f", fill_type="solid")
        complete = PatternFill(start_color="78d196", end_color="78d196", fill_type="solid")

        # Récurrence
        aucune = PatternFill(start_color="7b7f90", end_color="7b7f90", fill_type="solid")
        journalier = PatternFill(start_color="edbf72", end_color="edbf72", fill_type="solid")
        hebdomadaire = PatternFill(start_color="b1bee5", end_color="b1bee5", fill_type="solid")
        mensuelle = PatternFill(start_color="baabf2", end_color="baabf2", fill_type="solid")
        trimestrielle = PatternFill(start_color="94d4d0", end_color="94d4d0", fill_type="solid")
        annuelle = PatternFill(start_color="8caef6", end_color="8caef6", fill_type="solid")

        # Priorité
        faible = PatternFill(start_color="f3d659", end_color="f3d659", fill_type="solid")
        moyenne = PatternFill(start_color="edbf72", end_color="edbf72", fill_type="solid")
        haute = PatternFill(start_color="e68862", end_color="e68862", fill_type="solid")
        urgent = PatternFill(start_color="d0717f", end_color="d0717f", fill_type="solid")


        for i in range(len(data)):
            status = data[i]["column_values"]["Statut"]
            sheet["B" + str(i + 4)] = status
            if status == "À Faire":
                sheet["B" + str(i + 4)].fill = afaire
            elif status == "Planifié":
                sheet["B" + str(i + 4)].fill = planifie
            elif status == "En Cours":
                sheet["B" + str(i + 4)].fill = encours
            elif status == "Bloqué":
                sheet["B" + str(i + 4)].fill = bloque
            elif status == "Complété":
                sheet["B" + str(i + 4)].fill = complete

            sheet["C" + str(i + 4)] = data[i]["name"].capitalize()

            sheet["D" + str(i + 4)] = data[i]["column_values"]["Code Projet"]

            responsable = data[i]["column_values"]["Responsable"]
            responsable = responsable.split(".")[0]
            sheet["E" + str(i + 4)] = responsable.capitalize()

            sheet["F" + str(i + 4)] = data[i]["column_values"]["Date Limite"]

            recurrence = data[i]["column_values"]["Récurrence"]
            sheet["G" + str(i + 4)] = recurrence
            if recurrence == "Aucune":
                sheet["G" + str(i + 4)].fill = aucune
            elif recurrence == "Journalier":
                sheet["G" + str(i + 4)].fill = journalier
            elif recurrence == "Hebdomadaire":
                sheet["G" + str(i + 4)].fill = hebdomadaire
            elif recurrence == "Mensuelle":
                sheet["G" + str(i + 4)].fill = mensuelle
            elif recurrence == "Trimestrielle":
                sheet["G" + str(i + 4)].fill = trimestrielle
            elif recurrence == "Annuelle":
                sheet["G" + str(i + 4)].fill = annuelle

            priority = data[i]["column_values"]["Priorité"]
            sheet["H" + str(i + 4)] = priority
            if priority == "":
                pass
            elif priority == "Faible":
                sheet["H" + str(i + 4)].fill = faible
            elif priority == "Moyenne":
                sheet["H" + str(i + 4)].fill = moyenne
            elif priority == "Haute":
                sheet["H" + str(i + 4)].fill = haute
            elif priority == "Urgent":
                sheet["H" + str(i + 4)].fill = urgent

            sheet["I" + str(i + 4)] = data[i]["column_values"]["Temps Estimé"]
            sheet["J" + str(i + 4)] = data[i]["column_values"]["Département"]
            sheet["K" + str(i + 4)] = data[i]["column_values"]["Contacts"]
            sheet["L" + str(i + 4)] = data[i]["column_values"]["Téléphone"]
            sheet["M" + str(i + 4)] = data[i]["column_values"]["Lieu"]




        # Save the workbook
        workbook.save(workbook_path)
    
        return True
    else:
        tk.messagebox.showerror("Erreur", "Vous devez sélectionner un dossier pour enregistrer.")
    return False

# API Request
def call_api():
    # Headers for authentication
    headers = {
        'Authorization': API_KEY,
        'Content-Type': 'application/json'
    }

    items = []
    cursor = None  # Start with no cursor

    while True:
        # Dynamic variables for GraphQL query
        variables = {
            "boardId": [BOARD_ID],
            "cursor": cursor
        }

        # Make the API call
        response = requests.post(API_URL, headers=headers, json={'query': query, 'variables': variables})
        response_data = response.json()
        
        if response.status_code != 200 or 'errors' in response_data:
            print(f"Error fetching data: {response_data.get('errors')}")
            break

        # Extract items and cursor from the response
        board_data = response_data['data']['boards'][0]['items_page']
        items.extend(board_data['items'])
        cursor = board_data['cursor']
        
        # If cursor is None, it means there are no more pages to fetch
        if not cursor:
            break

    return items

# Function to reformat items into the desired structure (puts it into a dict)
def reformat_items(items):
    formatted_items = []
    for item in items:
        formatted_item = {
            "name": item['name'],
            "column_values": {cv['column']['title']: cv['text'] for cv in item['column_values']}
        }
        formatted_items.append(formatted_item)
    return formatted_items


root = ctk.CTk()
root.title("Monday à Excel")
root.geometry("700x500")
root.iconbitmap('./icon.ico')

# Choosing a font similar to the one in the image
design_font_name = "Segoe UI Semibold"

# Image Frame
imageFrame = ctk.CTkFrame(root, fg_color="red", width=300)
imageFrame.pack(side="left", fill="y")

# Disable Resize
root.resizable(False, False)

# Load the image
image = Image.open("./background.png")
image = image.resize((300, 500), Image.Resampling.LANCZOS)

# Convert the image to a CTkImage
ctk_image = ctk.CTkImage(light_image=image, size=(300, 500))

# Create a Label to display the image
image_label = ctk.CTkLabel(imageFrame, image=ctk_image, text="")
image_label.place(relwidth=1, relheight=1)  # Fill the entire frame

# Main Frame
main = ctk.CTkFrame(root)
main.pack(side="left", expand=True, fill="both")

# Use the chosen font for the bienvenue label
bienvenue = ctk.CTkLabel(master=main, text="Re-bienvenue!", font=(design_font_name, 44, "bold"), text_color="#047e7e")
bienvenue.pack(anchor="w", padx=(30,0), pady=(30,0))

mondayTitle = ctk.CTkLabel(master=main, text="Monday à Excel", font=(design_font_name, 20))
mondayTitle.pack(anchor="w", padx=(30,0))

# Key Frame
key = ctk.CTkFrame(main)
key.pack(pady=(10,0), padx=30, fill="x")

# API Section
apiFrame = ctk.CTkFrame(key, fg_color="transparent")

apiLabel = ctk.CTkLabel(master=apiFrame, text="API Key: ", font=(design_font_name, 14))
apiLabel.grid(row=0, column=0, padx=(10,5))

entry = ctk.CTkEntry(master=apiFrame, width=210, font=(design_font_name, 14))
entry.grid(row=0, column=1, padx=5)

saveBtn = ctk.CTkButton(master=apiFrame, text="Save", font=(design_font_name, 10), width=35, fg_color="#047e7e", command=update_api_key)
saveBtn.grid(row=0, column=2, padx=(5,10))

entry.insert(0, API_KEY)

apiFrame.pack(pady=(10,5))

# Board ID Section
boardFrame = ctk.CTkFrame(key, fg_color="transparent")

boardLabel = ctk.CTkLabel(boardFrame, text="Board ID: ", font=(design_font_name, 14))
boardLabel.grid(row=0, column=0, padx=(10,5))

boardEntry = ctk.CTkEntry(boardFrame, width=203, font=(design_font_name, 14))
boardEntry.grid(row=0, column=1, padx=5)

boardSaveBtn = ctk.CTkButton(boardFrame, text="Save", font=(design_font_name, 10), width=35, fg_color="#047e7e", command=update_board_id)
boardSaveBtn.grid(row=0, column=2, padx=(5,10))

boardEntry.insert(0, BOARD_ID)

boardFrame.pack(pady=(5,10))

# Ask Excel Section
excelFrame = ctk.CTkFrame(main)
excelFrame.pack(padx=30, pady=(10,5), fill="x")

excelFrame.columnconfigure(0, weight=1)
excelFrame.columnconfigure(1, weight=0)

switchExcel = ctk.CTkSwitch(master=excelFrame, text="Extraire Excel", font=(design_font_name, 14), command=toggle_switch_excel, progress_color="#047e7e")
switchExcel.grid(row=0, column=0, sticky="w", pady=10, padx=10)

# Folder Chooser Excel
chooseBtn = ctk.CTkButton(excelFrame, text="Choisir un dossier", font=(design_font_name, 12), command=lambda:choose_folder("excel"), fg_color="#047e7e")

excelPath = tk.Entry(root)

# Ask PDF Section
pdfFrame = ctk.CTkFrame(main)
pdfFrame.pack(padx=30, pady=(5,10), fill="x")

pdfFrame.columnconfigure(0, weight=1)
pdfFrame.columnconfigure(1, weight=0)

switchPDF = ctk.CTkSwitch(master=pdfFrame, text="Extraire Word", font=(design_font_name, 14), command=toggle_switch_pdf, progress_color="#047e7e")
switchPDF.grid(row=0, column=0, sticky="w", pady=10, padx=10)

chooseBtnPDF = ctk.CTkButton(pdfFrame, text="Choisir un dossier", font=(design_font_name, 12), command=lambda:choose_folder("pdf"), fg_color="#047e7e")

pdfPath = tk.Entry(root)

# Extract Button
exportBtn = ctk.CTkButton(master=main, text="Exporter", font=(design_font_name, 16), height=35, fg_color="#047e7e", command=export_info)
exportBtn.pack(pady=10)

#Texts of current path Frame
pathFrame = ctk.CTkFrame(master=main, fg_color="transparent")
pathFrame.pack(anchor="w", side="bottom", padx=10)

excelFilePath = ctk.CTkLabel(master=pathFrame, text="")
excelFilePath.grid(row=0, column=0, sticky="w")

pdfFilePath = ctk.CTkLabel(master=pathFrame, text="")
pdfFilePath.grid(row=1, column=0, sticky="w")


# Start the main loop
root.mainloop()